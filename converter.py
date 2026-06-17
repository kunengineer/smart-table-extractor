"""
converter.py – Trích xuất bảng từ PDF / Word / Ảnh sang Excel.
Không import streamlit. Dùng độc lập hoặc qua app.py.
"""

import os
import re
import logging
import platform
from typing import List, Optional

import pandas as pd
import numpy as np
import pdfplumber
import docx
import docx2txt
import pytesseract
import cv2
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from pdf2image import convert_from_path
from PIL import Image

# ══════════════════════════════════════════════════════════════════════════════
# LOGGING
# ══════════════════════════════════════════════════════════════════════════════
_FMT = "%(asctime)s [%(levelname)s] %(message)s"
logging.basicConfig(level=logging.INFO, format=_FMT)
logger = logging.getLogger("Table2Excel")


class _MemoryHandler(logging.Handler):
    """Ghi log vào bộ nhớ để Streamlit có thể đọc."""
    def __init__(self):
        super().__init__()
        self.records: List[str] = []

    def emit(self, record: logging.LogRecord) -> None:
        self.records.append(self.format(record))


_mem_handler = _MemoryHandler()
_mem_handler.setFormatter(logging.Formatter(_FMT))
logging.getLogger().addHandler(_mem_handler)


def get_captured_logs() -> List[str]:
    return list(_mem_handler.records)


def _clear_logs() -> None:
    _mem_handler.records.clear()


# ══════════════════════════════════════════════════════════════════════════════
# CUSTOM EXCEPTIONS
# ══════════════════════════════════════════════════════════════════════════════
class UnsupportedFormatError(Exception):
    pass


class ExtractionFailedError(Exception):
    pass


# ══════════════════════════════════════════════════════════════════════════════
# AUTO-CONFIG TESSERACT (Windows)
# ══════════════════════════════════════════════════════════════════════════════
if platform.system() == "Windows":
    _candidates = [
        r"C:\Program Files\Tesseract-OCR\tesseract.exe",
        r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
        os.path.join(
            os.environ.get("USERPROFILE", ""),
            r"AppData\Local\Programs\Tesseract-OCR\tesseract.exe",
        ),
    ]
    for _p in _candidates:
        if os.path.exists(_p):
            pytesseract.pytesseract.tesseract_cmd = _p
            break


# ══════════════════════════════════════════════════════════════════════════════
# HELPERS
# ══════════════════════════════════════════════════════════════════════════════
def _safe_str(v) -> str:
    """Chuyển giá trị bất kỳ sang string an toàn."""
    if v is None or (isinstance(v, float) and np.isnan(v)):
        return ""
    return str(v).strip()


def _rows_to_df(rows: List[List[str]]) -> Optional[pd.DataFrame]:
    """Tạo DataFrame từ list of rows, padding cột thiếu."""
    if len(rows) < 1:
        return None
    max_cols = max(len(r) for r in rows)
    if max_cols < 2:
        return None
    padded = [r + [""] * (max_cols - len(r)) for r in rows]
    return pd.DataFrame(padded)


def _get_clean_pdf_tables(page) -> List[List[List[str]]]:
    """
    Trích xuất bảng từ page pdfplumber bằng cách phân tích phân bố tọa độ ngang (Y)
    để lọc bỏ các nét thừa (như background rects hoặc cell borders lồng nhau)
    mà vẫn giữ nguyên cấu trúc dòng thực sự của bảng.
    """
    tables = page.find_tables()
    raw_tables = []
    if not tables:
        return []
        
    from collections import defaultdict
    
    for table in tables:
        tx0, ty0, tx1, ty1 = table.bbox
        t_width = tx1 - tx0
        
        # Thu thập các Y coordinate từ page.rects và page.lines trong vùng của table
        y_coords = []
        for r in page.rects:
            if r['top'] >= ty0 - 2 and r['bottom'] <= ty1 + 2:
                is_stroke = r.get('stroke', False) or r.get('linewidth', 0) > 0
                is_thin = r['height'] <= 5.0
                y_coords.append((r['top'], r['x0'], r['x1'], is_stroke, is_thin))
                y_coords.append((r['bottom'], r['x0'], r['x1'], is_stroke, is_thin))
                
        for l in page.lines:
            if l['top'] == l['bottom']:
                if l['top'] >= ty0 - 2 and l['top'] <= ty1 + 2:
                    y_coords.append((l['top'], l['x0'], l['x1'], True, True))
                    
        if not y_coords:
            # Nếu không tìm thấy tọa độ đường nào, trích xuất mặc định
            t_data = table.extract()
            if t_data:
                raw_tables.append(t_data)
            continue
            
        # Nhóm các Y coordinate cách nhau dưới 5.0 points
        y_groups = defaultdict(list)
        for y, x0, x1, stroke, thin in y_coords:
            found = False
            for gy in y_groups:
                if abs(y - gy) < 5.0:
                    y_groups[gy].append((y, x0, x1, stroke, thin))
                    found = True
                    break
            if not found:
                y_groups[y].append((y, x0, x1, stroke, thin))
                
        # Lọc các Y coordinate hợp lệ
        filtered_y = []
        for gy, items in y_groups.items():
            count = len(items)
            # Kiểm tra xem có đường dài (nét kẻ thực sự) hoặc số lượng cột đủ nhiều
            has_long_stroke_or_thin = any(
                (x[3] or x[4]) and (x[2] - x[1] >= 0.7 * t_width)
                for x in items
            )
            if count >= 10 or has_long_stroke_or_thin:
                avg_y = sum(x[0] for x in items) / len(items)
                filtered_y.append(avg_y)
                
        filtered_y.sort()
        
        # Nếu lọc được các đường ngang hợp lệ, trích xuất explicit
        if len(filtered_y) >= 2:
            settings = {
                'vertical_strategy': 'lines',
                'horizontal_strategy': 'explicit',
                'explicit_horizontal_lines': filtered_y,
            }
            t_data = page.extract_table(table_settings=settings)
            if t_data:
                raw_tables.append(t_data)
        else:
            t_data = table.extract()
            if t_data:
                raw_tables.append(t_data)
                
    return raw_tables


def _parse_text_lines(lines: List[str]) -> List[pd.DataFrame]:
    """
    Parse text thuần: tách theo tab hoặc 2+ khoảng trắng.
    Trả về danh sách DataFrame.
    """
    dfs: List[pd.DataFrame] = []
    buf: List[List[str]] = []

    def _flush():
        if len(buf) >= 2:
            df = _rows_to_df(buf)
            if df is not None:
                dfs.append(df)
        buf.clear()

    for raw in lines:
        line = raw.strip()
        if not line:
            _flush()
            continue
        if "\t" in line:
            fields = [f.strip() for f in line.split("\t")]
        else:
            fields = [f.strip() for f in re.split(r"\s{2,}", line)]

        if len(fields) >= 2:
            buf.append(fields)
        else:
            _flush()

    _flush()
    return dfs


# ══════════════════════════════════════════════════════════════════════════════
# OCR HELPER  (dùng chung cho PDF scan và image)
# ══════════════════════════════════════════════════════════════════════════════
# ══════════════════════════════════════════════════════════════════════════════
# OCR HELPER  (dùng chung cho PDF scan và image)
# ══════════════════════════════════════════════════════════════════════════════
def deskew_image(img_bgr: np.ndarray) -> np.ndarray:
    """
    Xác định góc nghiêng của ảnh dựa trên các đường thẳng phát hiện được
    và xoay ảnh để làm thẳng bảng giúp cải thiện chất lượng OCR và căn cột.
    """
    gray = cv2.cvtColor(img_bgr, cv2.COLOR_BGR2GRAY)
    thresh = cv2.adaptiveThreshold(
        gray, 255, cv2.ADAPTIVE_THRESH_MEAN_C, cv2.THRESH_BINARY_INV, 15, 4
    )
    
    # Phát hiện các nét ngang
    w_px = thresh.shape[1]
    h_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (max(w_px // 20, 20), 1))
    h_lines = cv2.morphologyEx(thresh, cv2.MORPH_OPEN, h_kernel, iterations=1)
    
    # Tìm các đoạn thẳng Hough
    lines = cv2.HoughLinesP(h_lines, 1, np.pi / 180, 50, minLineLength=w_px // 10, maxLineGap=20)
    
    if lines is None or len(lines) == 0:
        return img_bgr

    angles = []
    for line in lines:
        x1, y1, x2, y2 = line[0]
        angle = np.arctan2(y2 - y1, x2 - x1) * 180.0 / np.pi
        # Chỉ lấy các góc gần nằm ngang (trong khoảng -15 đến 15 độ)
        if -15 < angle < 15:
            angles.append(angle)
            
    if not angles:
        return img_bgr
        
    median_angle = np.median(angles)
    
    if abs(median_angle) < 0.2:
        return img_bgr # Quá thẳng rồi, không cần xoay
        
    logger.info(f"Phát hiện ảnh bị nghiêng {median_angle:.2f} độ. Đang tự động làm thẳng ảnh...")
    
    # Xoay ảnh xung quanh tâm
    (h, w) = img_bgr.shape[:2]
    center = (w // 2, h // 2)
    M = cv2.getRotationMatrix2D(center, median_angle, 1.0)
    # Dùng màu nền trắng cho phần khuyết khi xoay
    rotated = cv2.warpAffine(img_bgr, M, (w, h), flags=cv2.INTER_CUBIC, borderMode=cv2.BORDER_CONSTANT, borderValue=(255, 255, 255))
    return rotated


def crop_image_to_table(pil_img: Image.Image) -> Image.Image:
    """
    Sử dụng OpenCV để phát hiện vùng chứa bảng lớn nhất trong ảnh
    và cắt ảnh chỉ giữ lại vùng bảng này để lọc nhiễu tiêu đề/chữ ký bên ngoài.
    """
    # Chuyển PIL Image sang OpenCV BGR
    img_bgr = cv2.cvtColor(np.array(pil_img), cv2.COLOR_RGB2BGR)
    gray = cv2.cvtColor(img_bgr, cv2.COLOR_BGR2GRAY)
    thresh = cv2.adaptiveThreshold(
        gray, 255, cv2.ADAPTIVE_THRESH_MEAN_C, cv2.THRESH_BINARY_INV, 15, 4
    )

    h_px, w_px = thresh.shape
    h_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (max(w_px // 25, 20), 1))
    v_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (1, max(h_px // 25, 20)))
    h_lines = cv2.morphologyEx(thresh, cv2.MORPH_OPEN, h_kernel, iterations=2)
    v_lines = cv2.morphologyEx(thresh, cv2.MORPH_OPEN, v_kernel, iterations=2)
    grid = cv2.add(h_lines, v_lines)

    cnts, _ = cv2.findContours(grid, cv2.RETR_TREE, cv2.CHAIN_APPROX_SIMPLE)
    x_min, y_min, x_max, y_max = w_px, h_px, 0, 0
    found_table = False
    for cnt in cnts:
        x, y, w, h = cv2.boundingRect(cnt)
        # Bảng thường lớn hơn 25% chiều rộng và 5% chiều cao ảnh
        if w > w_px * 0.25 and h > h_px * 0.05:
            x_min = min(x_min, x)
            y_min = min(y_min, y)
            x_max = max(x_max, x + w)
            y_max = max(y_max, y + h)
            found_table = True

    if found_table:
        pad = 15
        x_min = max(0, x_min - pad)
        y_min = max(0, y_min - pad)
        x_max = min(w_px, x_max + pad)
        y_max = min(h_px, y_max + pad)
        logger.info(f"Đã phát hiện vùng bảng trong ảnh: x=({x_min}, {x_max}), y=({y_min}, {y_max}). Tiến hành cắt vùng bảng...")
        return pil_img.crop((x_min, y_min, x_max, y_max))
    
    logger.info("Không phát hiện vùng bảng rõ ràng. Giữ nguyên ảnh gốc.")
    return pil_img


def _ocr_to_df(pil_img: Image.Image) -> Optional[pd.DataFrame]:
    """
    Dùng pytesseract.image_to_data để ghép chữ thành bảng theo
    (block_num, line_num) → nhóm cột theo vị trí X.
    """
    # Cắt vùng bảng trước để tránh nhận diện các dòng tiêu đề và chữ ký bên ngoài
    cropped_img = crop_image_to_table(pil_img)
    
    try:
        data = pytesseract.image_to_data(
            cropped_img, output_type=pytesseract.Output.DICT,
            config="--psm 6"
        )
    except Exception as exc:
        logger.error(f"Tesseract lỗi: {exc}")
        raise ExtractionFailedError(
            "Không thể chạy Tesseract OCR. "
            "Kiểm tra lại cài đặt Tesseract và biến môi trường PATH."
        )

    words = []
    for i, txt in enumerate(data["text"]):
        txt = txt.strip()
        conf = int(data["conf"][i])
        if not txt or conf < 20:
            continue
        words.append({
            "text":  txt,
            "left":  data["left"][i],
            "top":   data["top"][i],
            "width": data["width"][i],
            "height": data["height"][i],
            "block": data["block_num"][i],
            "line":  data["line_num"][i],
        })

    if not words:
        return None

    from collections import defaultdict
    line_map = defaultdict(list)
    for w in words:
        line_map[(w["block"], w["line"])].append(w)

    lines_sorted = sorted(
        line_map.items(), key=lambda kv: np.mean([w["top"] for w in kv[1]])
    )

    img_w = cropped_img.width
    x_proj = np.zeros(img_w + 1, dtype=np.int32)
    for w in words:
        x0 = max(0, w["left"])
        x1 = min(img_w, w["left"] + w["width"])
        x_proj[x0:x1] += 1

    bands: List[List[int]] = []
    in_b = False
    b_start = 0
    for x in range(img_w + 1):
        if x_proj[x] > 0 and not in_b:
            b_start = x
            in_b = True
        elif x_proj[x] == 0 and in_b:
            bands.append([b_start, x])
            in_b = False

    merged: List[List[int]] = []
    # Ngưỡng gộp cột tự động dựa trên độ phân giải ảnh (1.5% chiều rộng vùng bảng)
    col_gap_threshold = max(8, int(img_w * 0.015))
    logger.info(f"Khoảng cách gộp cột tự động: {col_gap_threshold}px (chiều rộng bảng: {img_w}px)")
    for b in bands:
        if merged and b[0] - merged[-1][1] < col_gap_threshold:
            merged[-1][1] = b[1]
        else:
            merged.append(b)

    if not merged:
        return None

    grid: List[List[str]] = []
    for _, line_words in lines_sorted:
        row = [""] * len(merged)
        for w in sorted(line_words, key=lambda x: x["left"]):
            best, best_overlap = 0, -1
            for ci, (bx0, bx1) in enumerate(merged):
                overlap = min(w["left"] + w["width"], bx1) - max(w["left"], bx0)
                if overlap > best_overlap:
                    best_overlap = overlap
                    best = ci
            row[best] = (row[best] + " " + w["text"]).strip() if row[best] else w["text"]
        grid.append(row)

    return _rows_to_df(grid)


# ══════════════════════════════════════════════════════════════════════════════
# EXTRACTION
# ══════════════════════════════════════════════════════════════════════════════
def extract_from_pdf_text(file_path: str) -> List[pd.DataFrame]:
    """
    pdfplumber với chiến lược trích xuất và lọc đường kẻ ngang động.
    Tự động căn chỉnh và ghép các dòng bị tách.
    Fallback: tabula-py → OCR scan.
    """
    logger.info("Trích xuất PDF text bằng pdfplumber (custom horizontal lines strategy)...")
    dfs: List[pd.DataFrame] = []

    try:
        with pdfplumber.open(file_path) as pdf:
            for page_idx, page in enumerate(pdf.pages):
                raw_tables = _get_clean_pdf_tables(page)
                logger.info(f"Trang {page_idx+1}: pdfplumber tìm thấy {len(raw_tables)} bảng.")
                for raw in raw_tables:
                    if not raw or len(raw) < 2:
                        continue
                    df = _rows_to_df(raw)
                    if df is not None and not df.empty:
                        dfs.append(df)
        if dfs:
            logger.info(f"pdfplumber (lines) tìm thấy {len(dfs)} bảng hợp lệ.")
            return dfs
    except Exception as exc:
        logger.warning(f"pdfplumber lỗi: {exc}")

    # Fallback: tabula-py (lattice)
    logger.info("Thử tabula-py (lattice)...")
    try:
        import tabula  # noqa: PLC0415
        tabula_dfs = tabula.read_pdf(
            file_path, pages="all", multiple_tables=True,
            silent=True, lattice=True,
            encoding="utf-8"
        )
        for tdf in tabula_dfs or []:
            if isinstance(tdf, pd.DataFrame) and not tdf.empty:
                dfs.append(tdf)
        if dfs:
            logger.info(f"tabula-py (lattice) tìm thấy {len(dfs)} bảng.")
            return dfs
    except Exception as exc:
        logger.warning(f"tabula-py lattice lỗi: {exc}")

    # Fallback: tabula-py (stream)
    logger.info("Thử tabula-py (stream)...")
    try:
        import tabula  # noqa: PLC0415
        tabula_dfs = tabula.read_pdf(
            file_path, pages="all", multiple_tables=True,
            silent=True, stream=True,
            encoding="utf-8"
        )
        for tdf in tabula_dfs or []:
            if isinstance(tdf, pd.DataFrame) and not tdf.empty:
                dfs.append(tdf)
        if dfs:
            logger.info(f"tabula-py (stream) tìm thấy {len(dfs)} bảng.")
            return dfs
    except Exception as exc:
        logger.warning(f"tabula-py stream lỗi: {exc}")

    logger.info("Không tìm thấy bảng text → chuyển sang OCR scan...")
    return extract_from_pdf_scan(file_path)


def extract_from_pdf_scan(file_path: str) -> List[pd.DataFrame]:
    """pdf2image → OCR từng trang."""
    logger.info("Chuyển PDF scan thành ảnh (DPI=250)...")
    dfs: List[pd.DataFrame] = []

    # Tự động quét tìm poppler trên Windows
    poppler_path = None
    if platform.system() == "Windows":
        candidates = [
            r"C:\poppler\bin",
            r"C:\poppler\Library\bin",
            r"C:\Program Files\poppler\bin",
            r"C:\Program Files (x86)\poppler\bin",
        ]
        for p in candidates:
            if os.path.exists(p):
                poppler_path = p
                break
        
        # Nếu chưa tìm thấy, quét sơ bộ các thư mục con có tên poppler trong C:\
        if not poppler_path:
            try:
                for folder in os.listdir("C:\\"):
                    if "poppler" in folder.lower():
                        bin_path = os.path.join("C:\\", folder, "bin")
                        lib_bin = os.path.join("C:\\", folder, "Library", "bin")
                        if os.path.exists(lib_bin):
                            poppler_path = lib_bin
                            break
                        elif os.path.exists(bin_path):
                            poppler_path = bin_path
                            break
            except Exception:
                pass

    try:
        if poppler_path:
            logger.info(f"Đã phát hiện Poppler tại: {poppler_path}")
            images = convert_from_path(file_path, dpi=250, poppler_path=poppler_path)
        else:
            images = convert_from_path(file_path, dpi=250)
    except Exception as exc:
        logger.error(f"pdf2image lỗi: {exc}")
        raise ExtractionFailedError(
            "Không thể chuyển PDF sang ảnh. Vui lòng tải Poppler cho Windows "
            "(từ https://github.com/oschwartz10612/poppler-windows/releases), "
            "giải nén vào ổ đĩa C:\\ (ví dụ: C:\\poppler) và đảm bảo thư mục bin hoặc Library\\bin tồn tại."
        )

    for i, img in enumerate(images):
        logger.info(f"OCR trang {i + 1}/{len(images)}...")
        df = _ocr_to_df(img)
        if df is not None and not df.empty:
            dfs.append(df)

    if not dfs:
        raise ExtractionFailedError("Không trích xuất được bảng nào từ PDF scan.")
    return dfs


def extract_from_docx(file_path: str) -> List[pd.DataFrame]:
    """python-docx: đọc doc.tables trực tiếp → fallback parse paragraph text."""
    logger.info("Đọc bảng từ .docx...")
    dfs: List[pd.DataFrame] = []

    try:
        document = docx.Document(file_path)
        for tbl in document.tables:
            rows: List[List[str]] = []
            for row in tbl.rows:
                # Xử lý merged cells: lấy unique theo id của cell XML
                cells = list(row.cells)
                seen_ids: set = set()
                unique_cells: List[str] = []
                for c in cells:
                    cid = id(c._tc)
                    if cid not in seen_ids:
                        seen_ids.add(cid)
                        unique_cells.append(c.text.strip())
                rows.append(unique_cells)
            if len(rows) >= 2:
                df = _rows_to_df(rows)
                if df is not None:
                    dfs.append(df)
    except Exception as exc:
        logger.warning(f"Lỗi đọc .docx tables: {exc}")

    if not dfs:
        logger.info("Không có table trực tiếp → parse đoạn văn có cấu trúc bảng...")
        try:
            document = docx.Document(file_path)
            lines = [p.text for p in document.paragraphs]
            dfs = _parse_text_lines(lines)
        except Exception as exc:
            logger.warning(f"Lỗi parse paragraph: {exc}")

    return dfs


def extract_from_doc(file_path: str) -> List[pd.DataFrame]:
    """docx2txt → parse text."""
    logger.info("Đọc file .doc cũ bằng docx2txt...")
    try:
        text = docx2txt.process(file_path)
    except Exception as exc:
        raise ExtractionFailedError(
            f"Không đọc được .doc: {exc}\n"
            "Hãy mở file trong Word và lưu lại dưới dạng .docx."
        )

    if not text or not text.strip():
        raise ExtractionFailedError("File .doc rỗng hoặc không đọc được nội dung.")

    dfs = _parse_text_lines(text.split("\n"))
    if not dfs:
        raise ExtractionFailedError(
            "Không tìm thấy cấu trúc bảng trong .doc. "
            "Hãy chuyển sang .docx để nhận diện tốt hơn."
        )
    return dfs


def extract_from_image(file_path: str) -> List[pd.DataFrame]:
    """
    OpenCV phát hiện đường kẻ → crop ô → OCR.
    Fallback: OCR toàn ảnh bằng _ocr_to_df.
    """
    logger.info("Phát hiện đường kẻ bảng bằng OpenCV...")
    img_bgr_raw = cv2.imread(file_path)
    if img_bgr_raw is None:
        raise ExtractionFailedError(f"Không đọc được file ảnh: {file_path}")

    # 1. Tự động xoay làm thẳng ảnh nếu bị nghiêng
    img_bgr = deskew_image(img_bgr_raw)

    # Chuyển đổi sang PIL Image để tiến hành cắt vùng chứa bảng
    pil_img = Image.fromarray(cv2.cvtColor(img_bgr, cv2.COLOR_BGR2RGB))
    pil_img_cropped = crop_image_to_table(pil_img)

    # Chuyển ngược lại sang OpenCV để chạy thuật toán tìm ô lưới
    img_bgr_cropped = cv2.cvtColor(np.array(pil_img_cropped), cv2.COLOR_RGB2BGR)

    gray = cv2.cvtColor(img_bgr_cropped, cv2.COLOR_BGR2GRAY)
    thresh = cv2.adaptiveThreshold(
        gray, 255, cv2.ADAPTIVE_THRESH_MEAN_C, cv2.THRESH_BINARY_INV, 15, 4
    )

    h_px, w_px = thresh.shape
    h_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (max(w_px // 25, 20), 1))
    v_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (1, max(h_px // 25, 20)))
    h_lines = cv2.morphologyEx(thresh, cv2.MORPH_OPEN, h_kernel, iterations=2)
    v_lines = cv2.morphologyEx(thresh, cv2.MORPH_OPEN, v_kernel, iterations=2)
    grid = cv2.add(h_lines, v_lines)

    cnts, hier = cv2.findContours(grid, cv2.RETR_CCOMP, cv2.CHAIN_APPROX_SIMPLE)

    cells = []
    if hier is not None:
        for i, cnt in enumerate(cnts):
            parent = hier[0][i][3]
            if parent == -1:
                continue
            x, y, cw, ch = cv2.boundingRect(cnt)
            if cw > 15 and ch > 10 and cw < w_px * 0.95:
                cells.append((x, y, cw, ch))

    if len(cells) < 4:
        logger.info("Không đủ ô lưới trong vùng bảng → fallback OCR vùng bảng...")
        df = _ocr_to_df(pil_img_cropped)
        return [df] if df is not None and not df.empty else []

    logger.info(f"OpenCV phát hiện {len(cells)} ô trong vùng bảng. Tiến hành OCR từng ô...")

    cells.sort(key=lambda c: c[1])
    row_groups: List[List[tuple]] = []
    cur: List[tuple] = [cells[0]]
    for cell in cells[1:]:
        avg_y = np.mean([c[1] for c in cur])
        avg_h = np.mean([c[3] for c in cur])
        if abs(cell[1] - avg_y) < avg_h * 0.6:
            cur.append(cell)
        else:
            row_groups.append(sorted(cur, key=lambda c: c[0]))
            cur = [cell]
    row_groups.append(sorted(cur, key=lambda c: c[0]))

    all_x = [(c[0], c[0] + c[2]) for c in cells]
    all_x.sort()
    col_bands: List[List[int]] = []
    for x0, x1 in all_x:
        if col_bands and x0 <= col_bands[-1][1] + 5:
            col_bands[-1][1] = max(col_bands[-1][1], x1)
        else:
            col_bands.append([x0, x1])
    col_bands.sort(key=lambda b: b[0])

    grid_data: List[List[str]] = []
    for r_cells in row_groups:
        row_data = [""] * len(col_bands)
        for (cx, cy, cw, ch) in r_cells:
            best_ci, best_ov = 0, -1
            for ci, (bx0, bx1) in enumerate(col_bands):
                ov = min(cx + cw, bx1) - max(cx, bx0)
                if ov > best_ov:
                    best_ov = ov
                    best_ci = ci
            pad = 3
            crop = pil_img_cropped.crop((
                max(0, cx + pad), max(0, cy + pad),
                min(w_px, cx + cw - pad), min(h_px, cy + ch - pad),
            ))
            try:
                txt = pytesseract.image_to_string(crop, config="--psm 7").strip()
            except Exception as exc:
                logger.error(f"Tesseract OCR lỗi: {exc}")
                raise ExtractionFailedError(
                    "Không thể chạy Tesseract OCR. Vui lòng tải và cài đặt Tesseract OCR "
                    "cho Windows (ví dụ từ github.com/UB-Mannheim/tesseract/wiki) và cài đặt vào "
                    "thư mục mặc định (C:\\Program Files\\Tesseract-OCR). Chương trình sẽ tự động nhận diện."
                )
            row_data[best_ci] = txt
        grid_data.append(row_data)

    df = _rows_to_df(grid_data)
    # Nếu bảng thu được có ít hơn 4 cột (thường là lỗi gộp cột khi thiếu nét dọc),
    # ta fallback sang OCR chiếu dòng để thu được bảng có số lượng cột đúng hơn.
    if df is not None and not df.empty and len(df.columns) >= 4:
        return [df]

    logger.info("Bảng từ ô lưới có ít hơn 4 cột hoặc rỗng. Fallback sang OCR chiếu dòng...")
    df2 = _ocr_to_df(pil_img_cropped)
    return [df2] if df2 is not None and not df2.empty else []


# ══════════════════════════════════════════════════════════════════════════════
# POST-PROCESSING
# ══════════════════════════════════════════════════════════════════════════════
def is_layout_table(df: pd.DataFrame) -> bool:
    """
    Kiểm tra xem DataFrame có phải là bảng bố cục tiêu đề thư (letterhead)
    hoặc bảng chữ ký/footer hay không.
    """
    if df is None or df.empty:
        return True
        
    # Bảng bố cục thư hoặc chữ ký thường có rất ít dòng dữ liệu
    if len(df) <= 2:
        # Lấy tất cả nội dung bao gồm cả tên cột (headers) và dữ liệu
        all_elements = list(df.columns.astype(str)) + list(df.astype(str).values.flatten())
        text = " ".join(all_elements).lower()
        header_kws = {"cộng hòa xã hội", "độc lập – tự do", "độc lập - tự do", "độc lập tự do", "hạnh phúc"}
        footer_kws = {"trưởng đơn vị", "người đề nghị", "người đề xuất", "người lập", "kế toán trưởng", "giám đốc", "phòng hcns"}
        if any(kw in text for kw in header_kws | footer_kws):
            return True
            
    return False


def crop_table_header_footer(df: pd.DataFrame) -> pd.DataFrame:
    """
    Cắt bỏ các dòng tiêu đề văn bản bên ngoài bảng (header) ở đầu 
    và các thông tin chữ ký bên ngoài bảng (footer) ở cuối.
    Chỉ giữ lại phần bảng dữ liệu thực sự.
    """
    if df is None or df.empty:
        return df

    # Chuyển tất cả sang string và chuẩn hóa để tìm kiếm từ khóa
    rows_str = [[str(v).strip().lower() for v in row] for row in df.values]

    header_kws = {
        "stt", "loại", "tên", "sản phẩm", "đơn vị", "số lượng", "đơn giá", 
        "thành tiền", "mục đích", "sử dụng", "thời gian", "ghi chú", 
        "ngày", "mã", "ký", "hạn", "mặt hàng"
    }
    
    footer_kws = {
        "trưởng phòng", "giám đốc", "người lập", "người đề nghị", 
        "người đề xuất", "kế toán", "ngày...tháng", "xác nhận", 
        "phòng hcns", "trưởng đơn vị", "ký tên", "ký rõ"
    }

    # 1. Tìm dòng tiêu đề thực sự của bảng (chứa nhiều từ khóa cột nhất)
    best_header_idx = 0
    max_kws = 0
    
    for idx, row in enumerate(rows_str):
        kws_count = sum(1 for cell in row if any(kw in cell for kw in header_kws))
        if kws_count > max_kws:
            max_kws = kws_count
            best_header_idx = idx

    # Cắt phần đầu bảng từ dòng header trở đi
    if max_kws >= 2:
        df = df.iloc[best_header_idx:].reset_index(drop=True)
        rows_str = [[str(v).strip().lower() for v in row] for row in df.values]

    # 2. Cắt bỏ footer từ phía cuối bảng lên
    last_idx = len(df) - 1
    while last_idx >= 0:
        row = rows_str[last_idx]
        if all(cell == "" or cell == "nan" for cell in row):
            last_idx -= 1
            continue
        is_f = any(any(kw in cell for kw in footer_kws) for cell in row)
        if is_f:
            last_idx -= 1
        else:
            break

    return df.iloc[:last_idx + 1].reset_index(drop=True)


def post_process(df: pd.DataFrame) -> pd.DataFrame:
    """
    Chuẩn hóa: strip, xóa dòng/cột rỗng hoàn toàn,
    cắt tiêu đề/chữ ký dư thừa, và dùng dòng đầu làm header.
    """
    if df is None or df.empty:
        return pd.DataFrame()

    # Cắt bỏ header/footer dư thừa trước tiên
    df = crop_table_header_footer(df)

    def _strip_cell(v):
        return v.strip() if isinstance(v, str) else v

    try:
        df = df.map(_strip_cell)
    except AttributeError:
        df = df.applymap(_strip_cell)

    df = df.fillna("")

    # Xóa dòng toàn rỗng
    df = df[~df.apply(lambda row: all(str(v).strip() == "" for v in row), axis=1)]
    # Xóa cột toàn rỗng
    df = df.loc[:, ~df.apply(lambda col: all(str(v).strip() == "" for v in col), axis=0)]
    df = df.reset_index(drop=True)

    if df.empty or df.shape[1] < 2:
        return pd.DataFrame()

    # Kiểm tra mật độ dữ liệu tối thiểu (5% - rất thoáng)
    total = df.size
    filled = df.astype(str).apply(lambda col: col.str.strip() != "").sum().sum()
    if total > 0 and filled / total < 0.05:
        logger.info(f"Bảng bị bỏ qua: mật độ dữ liệu chỉ {filled/total:.1%} < 5%.")
        return pd.DataFrame()

    # Dùng dòng đầu làm header
    if len(df) >= 2:
        first = df.iloc[0].astype(str)
        has_text = any(v.strip() for v in first)
        if has_text:
            raw_header = list(first)
            seen: dict = {}
            clean: List[str] = []
            for i, h in enumerate(raw_header):
                h = h.strip() or f"Cột_{i + 1}"
                if h in seen:
                    seen[h] += 1
                    clean.append(f"{h}_{seen[h]}")
                else:
                    seen[h] = 0
                    clean.append(h)
            df.columns = clean
            df = df.iloc[1:].reset_index(drop=True)
            logger.info("Đã đặt dòng đầu làm header cột.")

    # --- Lọc bỏ dòng trống (chỉ có STT hoặc rỗng toàn bộ) và dòng Tổng cộng ---
    if not df.empty:
        # Xóa các cột liên quan đến Tên file nguồn
        cols_to_drop = [
            c for c in df.columns 
            if any(k in str(c).lower() for k in ["file nguồn", "file nguon", "tên file", "ten file", "source file"])
        ]
        if cols_to_drop:
            df = df.drop(columns=cols_to_drop)
            logger.info(f"Đã xóa các cột tên file nguồn: {cols_to_drop}")

    if not df.empty:
        # 1. Tìm cột STT bằng header hoặc dữ liệu số thứ tự tăng dần
        stt_col = None
        stt_kws = ["stt", "no.", "index", "số thứ tự", "sothutu"]
        for c in df.columns:
            c_norm = str(c).strip().lower().replace('\n', ' ').replace('\r', ' ')
            if any(kw in c_norm for kw in stt_kws):
                stt_col = c
                break
                
        if stt_col is None:
            for c in df.columns:
                col_vals = df[c].astype(str).str.strip().tolist()
                non_empty = [v for v in col_vals if v not in ("", "nan", "none")]
                if len(non_empty) >= 2:
                    try:
                        ints = [int(float(v)) for v in non_empty]
                        if all(0 < x < 500 for x in ints) and ints == sorted(ints):
                            stt_col = c
                            break
                    except ValueError:
                        continue
                
        # 2. Định nghĩa dòng trống (các cột ngoại trừ STT đều rỗng hoặc bằng 0)
        check_cols = [c for c in df.columns if c != stt_col]
        if check_cols:
            def is_empty_data_row(row):
                for c in check_cols:
                    val = str(row[c]).replace('\xa0', ' ').strip().lower()
                    if val not in ("", "nan", "none", "0", "0.0"):
                        return False
                return True
            df = df[~df.apply(is_empty_data_row, axis=1)]
            
        # 3. Định nghĩa dòng Tổng cộng với chuẩn hóa Unicode (NFC)
        import unicodedata
        total_kws = {"tổng cộng", "tong cong", "tổng số", "tong so", "cộng:", "cong:", "cộng dòng", "cong dong", "cộng", "total", "grand total"}
        total_kws_nfc = {unicodedata.normalize('NFC', kw) for kw in total_kws}
        
        def is_total_row(row):
            for cell in row:
                cell_val = str(cell).replace('\xa0', ' ').strip().lower().replace(':', '').strip()
                cell_norm = unicodedata.normalize('NFC', cell_val)
                if any(kw == cell_norm or cell_norm.startswith(kw) for kw in total_kws_nfc) or "t픀" in cell_norm or "c픀" in cell_norm:
                    return True
            return False
        df = df[~df.apply(is_total_row, axis=1)]
        df = df.reset_index(drop=True)

    return df


# ══════════════════════════════════════════════════════════════════════════════
# SAVE TO EXCEL
# ══════════════════════════════════════════════════════════════════════════════
def save_to_excel(dataframes: List[pd.DataFrame], output_path: str, sheet_names: Optional[List[str]] = None) -> str:
    if not dataframes:
        raise ValueError("Không có bảng nào để lưu.")

    wb = openpyxl.Workbook()
    wb.remove(wb.active)

    # Palette màu xanh dương đậm hiện đại, dịu mắt
    hdr_fill = PatternFill("solid", fgColor="1F4E79")
    hdr_font = Font(name="Segoe UI", bold=True, size=11, color="FFFFFF")
    zebra    = PatternFill("solid", fgColor="F2F7FA")
    body_fnt = Font(name="Segoe UI", size=10)
    
    # Border mỏng màu xám nhạt tinh tế
    thin_border = Border(
        left=Side(style='thin', color='D9D9D9'),
        right=Side(style='thin', color='D9D9D9'),
        top=Side(style='thin', color='D9D9D9'),
        bottom=Side(style='thin', color='D9D9D9')
    )
    
    center_al = Alignment(horizontal="center", vertical="center", wrap_text=True)
    left_al   = Alignment(horizontal="left",   vertical="center", wrap_text=True)
    right_al  = Alignment(horizontal="right",  vertical="center", wrap_text=True)

    for idx, df in enumerate(dataframes):
        title = f"Bảng_{idx + 1}"
        if sheet_names and idx < len(sheet_names):
            title = sheet_names[idx]
        ws = wb.create_sheet(title=title)
        
        # Luôn hiển thị đường lưới trong Excel
        ws.sheet_view.showGridLines = True
        
        headers = list(df.columns)

        # Ghi Header
        ws.append(headers)
        for ci in range(1, len(headers) + 1):
            c = ws.cell(1, ci)
            c.fill = hdr_fill
            c.font = hdr_font
            c.alignment = center_al
            c.border = thin_border
        ws.row_dimensions[1].height = 28

        # Ghi các dòng dữ liệu
        for ri, row_vals in enumerate(df.values, start=2):
            native_row = []
            for v in row_vals:
                if v is None or (isinstance(v, float) and np.isnan(v)):
                    native_row.append(None)
                else:
                    val_str = str(v).strip()
                    # Loại bỏ đơn vị tiền tệ và dấu phẩy phân tách phần nghìn để parse sang số chính xác
                    val_clean = val_str.lower().replace("đ", "").replace("vnd", "").replace("vân", "").strip()
                    val_clean = val_clean.replace(",", "")
                    try:
                        if val_clean.endswith(".0"):
                            native_row.append(int(val_clean[:-2]))
                        elif val_clean.isdigit():
                            native_row.append(int(val_clean))
                        else:
                            val_float = float(val_clean)
                            if val_float.is_integer():
                                native_row.append(int(val_float))
                            else:
                                native_row.append(val_float)
                    except ValueError:
                        native_row.append(v.item() if hasattr(v, "item") else v)

            ws.append(native_row)
            fill = zebra if ri % 2 == 0 else None
            ws.row_dimensions[ri].height = 20
            
            for ci in range(1, len(headers) + 1):
                cell = ws.cell(ri, ci)
                cell.font = body_fnt
                cell.border = thin_border
                if fill:
                    cell.fill = fill
                
                col_name = str(headers[ci-1]).lower()
                val = cell.value
                
                # Căn lề thông minh dựa trên kiểu dữ liệu và cột
                if col_name == "stt" or (val is not None and isinstance(val, int) and val < 1000 and col_name == "stt"):
                    cell.alignment = center_al
                elif isinstance(val, (int, float)):
                    cell.alignment = right_al
                    # Định dạng hiển thị số sạch sẽ
                    if isinstance(val, int):
                        cell.number_format = "#,##0"
                    else:
                        cell.number_format = "#,##0.00"
                else:
                    cell.alignment = left_al

        # Thêm dòng Tổng cộng nếu có cột Số lượng hoặc Thành tiền
        qty_col_idx = None
        total_col_idx = None
        
        qty_kws = ["số lượng", "so luong", "số lượng", "qty", "quantity", "sl"]
        total_kws = ["thành tiền", "thanh tien", "thành tiền", "amount", "total", "tt"]
        
        for ci, h in enumerate(headers, 1):
            h_norm = str(h).strip().lower().replace('\n', ' ').replace('\r', ' ')
            if "stt" in h_norm:
                continue
            if any(kw in h_norm for kw in qty_kws) and qty_col_idx is None:
                qty_col_idx = ci
            elif any(kw in h_norm for kw in total_kws) and total_col_idx is None:
                total_col_idx = ci

        if len(df) > 0 and (qty_col_idx or total_col_idx):
            summary_row_idx = ws.max_row + 1
            ws.row_dimensions[summary_row_idx].height = 24
            
            # Đặt nhãn "Tổng cộng :" vào cột ngay trước cột tổng đầu tiên (hoặc cột 1)
            label_col = 1
            if qty_col_idx and qty_col_idx > 1:
                label_col = qty_col_idx - 1
            elif total_col_idx and total_col_idx > 1:
                label_col = total_col_idx - 1
                
            # Đảm bảo label_col không trùng với bất kỳ cột tổng nào
            if label_col == qty_col_idx or label_col == total_col_idx:
                label_col = 1
                
            ws.cell(summary_row_idx, label_col, "Tổng cộng :")
            lbl_cell = ws.cell(summary_row_idx, label_col)
            lbl_cell.font = Font(name="Segoe UI", bold=True, size=10)
            lbl_cell.alignment = Alignment(horizontal="right", vertical="center")
            
            # Áp dụng viền và font bold cho toàn bộ hàng Tổng cộng
            for ci in range(1, len(headers) + 1):
                c = ws.cell(summary_row_idx, ci)
                c.border = thin_border
                c.font = Font(name="Segoe UI", bold=True, size=10)
                
                if ci == qty_col_idx:
                    col_letter = get_column_letter(ci)
                    c.value = f"=SUM({col_letter}2:{col_letter}{summary_row_idx - 1})"
                    c.alignment = right_al
                    c.number_format = "#,##0"
                elif ci == total_col_idx:
                    col_letter = get_column_letter(ci)
                    c.value = f"=SUM({col_letter}2:{col_letter}{summary_row_idx - 1})"
                    c.alignment = right_al
                    c.number_format = "#,##0.00"

        ws.freeze_panes = "A2"

        # Tự động căn chỉnh độ rộng cột với biên rộng hơn tí cho dễ nhìn
        for ci, col_name in enumerate(headers, 1):
            col_letter = get_column_letter(ci)
            max_len = len(str(col_name))
            for ri in range(2, ws.max_row + 1):
                val = ws.cell(ri, ci).value
                if val is not None:
                    lines = str(val).split("\n")
                    max_len = max(max_len, max(len(ln) for ln in lines))
            ws.column_dimensions[col_letter].width = min(max_len + 4, 60)

    wb.save(output_path)
    total_rows = sum(len(df) for df in dataframes)
    logger.info(
        f"Lưu thành công: {output_path} "
        f"({len(dataframes)} sheet, {total_rows} dòng)"
    )
    return output_path


# ══════════════════════════════════════════════════════════════════════════════
# MAIN ORCHESTRATOR
# ══════════════════════════════════════════════════════════════════════════════
def detect_and_convert(file_path: str, output_path: Optional[str] = None) -> str:
    """
    Phát hiện loại file → trích xuất bảng → post-process → lưu Excel.
    Trả về đường dẫn file Excel đã lưu.
    """
    _clear_logs()

    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Không tìm thấy file: {file_path}")

    if output_path is None:
        base = os.path.splitext(file_path)[0]
        output_path = f"{base}_converted.xlsx"

    logger.info(f"Bắt đầu xử lý: {os.path.basename(file_path)}")

    # ── Phát hiện loại file ──────────────────────────────────────────────────
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        try:
            with pdfplumber.open(file_path) as pdf:
                total_text = "".join([p.extract_text() or "" for p in pdf.pages])
            file_type = "pdf_text" if len(total_text.strip()) > 50 else "pdf_scan"
        except Exception:
            file_type = "pdf_scan"
    elif ext == ".docx":
        file_type = "docx"
    elif ext == ".doc":
        file_type = "doc"
    elif ext in {".png", ".jpg", ".jpeg", ".tiff", ".bmp", ".webp"}:
        file_type = "image"
    else:
        raise UnsupportedFormatError(
            f"Định dạng '{ext}' không được hỗ trợ.\n"
            "Dùng: .pdf, .docx, .doc, .png, .jpg, .jpeg"
        )
    logger.info(f"Loại file: {file_type}")

    # ── Trích xuất ───────────────────────────────────────────────────────────
    dispatch = {
        "pdf_text": extract_from_pdf_text,
        "pdf_scan": extract_from_pdf_scan,
        "docx":     extract_from_docx,
        "doc":      extract_from_doc,
        "image":    extract_from_image,
    }
    raw_dfs = dispatch[file_type](file_path)

    # ── Post-process ─────────────────────────────────────────────────────────
    processed = []
    for i, df in enumerate(raw_dfs):
        logger.info(f"Post-process bảng {i + 1} (shape={df.shape})...")
        clean = post_process(df)
        if not clean.empty and not is_layout_table(clean):
            processed.append(clean)
            logger.info(f"  → Hợp lệ: {clean.shape[0]} dòng × {clean.shape[1]} cột.")
        else:
            logger.info(f"  → Bỏ qua (rỗng hoặc bảng bố cục sau khi lọc).")

    if not processed:
        raise ExtractionFailedError(
            "Không trích xuất được bảng hợp lệ nào từ file.\n"
            "Kiểm tra: file có chứa bảng dữ liệu không? Ảnh/scan có đủ nét không?"
        )

    logger.info(f"Tổng {len(processed)} bảng hợp lệ → đang lưu Excel...")
    return save_to_excel(processed, output_path)